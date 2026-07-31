from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
MARKDOWN_PATH = HERE / "bao_cao_tuan_5.md"
OUTPUT_PATH = HERE / "bao_cao_tuan_5.docx"
ASSET_DIR = HERE.parents[2] / "output" / "week5" / "report_assets"
ARCHITECTURE_PATH = ASSET_DIR / "week5_architecture.png"

NAVY = RGBColor(25, 55, 90)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
LIGHT = "F2F4F7"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
CONTENT_DXA = 9360


def set_font(run, size=11, color=BLACK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("DE GENESIS  |  BÁO CÁO TUẦN 5"), 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("Trang "), 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_cover(document: Document) -> None:
    for _ in range(5):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_font(p.add_run("DE GENESIS  |  DATA ENGINEERING LAB"), 11, BLUE, True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("BÁO CÁO TUẦN 5"), 30, NAVY, True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_font(p.add_run("AIRFLOW, NIFI VÀ TÍCH HỢP\nPROMOTION API"), 15.5, DARK_BLUE, True)
    metadata = [
        ("Dự án", "DE Genesis - Lộ trình thực hành Data Engineering"),
        ("Phạm vi", "Hai pipeline orchestration và API integration"),
        ("Công nghệ", "FastAPI, PostgreSQL, Airflow 2.9.3, NiFi 1.27.0, Spark 3.5.1"),
        ("Ngày báo cáo", "24/07/2026"),
    ]
    for label, value in metadata:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(f"{label}: "), 10.5, MUTED, True)
        set_font(p.add_run(value), 10.5)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(46)
    set_font(p.add_run("Báo cáo dựa trên mã nguồn và bằng chứng kiểm thử trong repo."), 9.5, MUTED, italic=True)
    document.add_page_break()


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.is_file():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_box(draw, rect, title, subtitle, fill):
    draw.rounded_rectangle(rect, radius=18, fill=fill, outline="#2E74B5", width=3)
    x1, y1, x2, _ = rect
    title_font = font(26, True)
    sub_font = font(21)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((x1 + x2 - title_box[2]) / 2, y1 + 36), title, font=title_font, fill="#19375A")
    lines = subtitle.split("\n")
    for i, line in enumerate(lines):
        box = draw.textbbox((0, 0), line, font=sub_font)
        draw.text(((x1 + x2 - box[2]) / 2, y1 + 88 + i * 30), line, font=sub_font, fill="#5A636E")


def arrow(draw, start, end):
    draw.line([start, end], fill="#2E74B5", width=5)
    draw.polygon([(end[0], end[1]), (end[0] - 16, end[1] - 10), (end[0] - 16, end[1] + 10)], fill="#2E74B5")


def generate_architecture() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 850), "white")
    draw = ImageDraw.Draw(image)
    title = "KIẾN TRÚC SO SÁNH HAI PIPELINE TUẦN 5"
    box = draw.textbbox((0, 0), title, font=font(38, True))
    draw.text(((1800 - box[2]) / 2, 28), title, font=font(38, True), fill="#19375A")
    draw_box(draw, (60, 315, 340, 535), "Promotion API", "250 records\nFailure scenarios", "#E8EEF5")
    top = [
        (430, 120, 710, 300, "Airflow", "API ingestion\nRetry / pagination"),
        (790, 120, 1070, 300, "Raw Airflow", "PostgreSQL\nBatch + hash"),
        (1150, 120, 1430, 300, "Spark core", "Join Olist\nTính discount"),
        (1510, 120, 1770, 300, "Curated", "Airflow result\nQuality report"),
    ]
    bottom = [
        (430, 560, 710, 740, "NiFi", "InvokeHTTP\nRoute / validate"),
        (790, 560, 1070, 740, "Raw NiFi", "PostgreSQL\nBatch + hash"),
        (1150, 560, 1430, 740, "Airflow REST", "Trigger DAG\nIdempotent run"),
        (1510, 560, 1770, 740, "Spark core", "Shared logic\nCurated result"),
    ]
    for items in (top, bottom):
        for i, (*rect, title_text, subtitle) in enumerate(items):
            draw_box(draw, tuple(rect), title_text, subtitle, "#F2F4F7")
            if i < len(items) - 1:
                arrow(draw, (rect[2] + 10, (rect[1] + rect[3]) // 2), (items[i + 1][0] - 10, (rect[1] + rect[3]) // 2))
    arrow(draw, (340, 390), (430, 210))
    arrow(draw, (340, 460), (430, 650))
    image.save(ARCHITECTURE_PATH)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            set_font(paragraph.add_run(part[1:-1]), 9.5, DARK_BLUE, name="Consolas")
        elif part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), 11, bold=True)
        else:
            set_font(paragraph.add_run(part), 11)


def add_table(document, rows):
    columns = len(rows[0])
    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [1800, 3300, 4260]
    else:
        widths = [CONTENT_DXA // columns] * columns
        widths[-1] += CONTENT_DXA - sum(widths)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(value), 9.5, WHITE if row_index == 0 else BLACK, row_index == 0)
            if row_index == 0:
                shade_cell(cell, "2E74B5")
    set_table_geometry(table, widths)
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def parse_body(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("# 1."))
    i = start
    major_count = 0
    paragraph_buffer = []

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = document.add_paragraph()
            inline(p, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        if line == "[[ARCHITECTURE]]":
            flush()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            picture_run = p.add_run()
            picture_run.add_picture(str(ARCHITECTURE_PATH), width=Inches(6.35))
            doc_property = picture_run._r.xpath(".//wp:docPr")[0]
            doc_property.set(
                "descr",
                "Sơ đồ so sánh pipeline Airflow-centric và NiFi-centric dùng chung Promotion API và Spark core.",
            )
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(8)
            set_font(caption.add_run("Hình 1. Kiến trúc hai pipeline tuần 5"), 9, MUTED, italic=True)
        elif line.startswith("# "):
            flush()
            major_count += 1
            if major_count > 1 and major_count in {4, 7, 10, 12, 14}:
                document.add_page_break()
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            flush()
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            flush()
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            flush()
            p = document.add_paragraph(style="List Bullet")
            inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            flush()
            p = document.add_paragraph(style="List Number")
            inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("|"):
            flush()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|\s*:?-+", row)
            ]
            add_table(document, parsed)
            continue
        elif not line:
            flush()
        else:
            paragraph_buffer.append(line)
        i += 1
    flush()


def build() -> None:
    generate_architecture()
    document = Document()
    configure(document)
    add_cover(document)
    parse_body(document, MARKDOWN_PATH.read_text(encoding="utf-8"))
    document.core_properties.title = "Báo cáo thực hành tuần 5 - DE Genesis"
    document.core_properties.subject = "Airflow, NiFi và tích hợp Promotion API"
    document.core_properties.author = "DE Genesis"
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
