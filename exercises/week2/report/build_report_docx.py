from __future__ import annotations

import importlib.util
import os
import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


HERE = Path(__file__).resolve().parent
PROJECT_ROOT = HERE.parents[2]
MARKDOWN_PATH = HERE / "bao_cao_tuan_2.md"
OUTPUT_PATH = HERE / "bao_cao_tuan_2.docx"
ASSET_DIR = HERE / "_generated_assets"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = RGBColor(31, 78, 121)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
VERY_LIGHT = "F7F9FC"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


class FallbackTableGeometry:
    """Geometry OOXML tối thiểu khi máy chạy không cài Documents skill."""

    @staticmethod
    def column_widths_from_weights(
        weights: list[float], total_width_dxa: int
    ) -> list[int]:
        if not weights or any(weight <= 0 for weight in weights):
            raise ValueError("Trọng số cột phải là các số dương")
        total_weight = float(sum(weights))
        widths = [
            int(round(total_width_dxa * weight / total_weight)) for weight in weights
        ]
        widths[-1] += total_width_dxa - sum(widths)
        if any(width <= 0 for width in widths):
            raise ValueError(f"Độ rộng cột không hợp lệ: {widths}")
        return widths

    @staticmethod
    def _ensure_child(parent, tag: str):
        child = parent.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            parent.append(child)
        return child

    @classmethod
    def _set_width(cls, parent, tag: str, width_dxa: int) -> None:
        width = cls._ensure_child(parent, tag)
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(int(width_dxa)))

    @classmethod
    def apply_table_geometry(
        cls,
        table,
        column_widths_dxa: list[int],
        *,
        table_width_dxa: int,
        indent_dxa: int,
        cell_margins_dxa: dict[str, int],
    ) -> None:
        widths = [int(width) for width in column_widths_dxa]
        if not widths or any(width <= 0 for width in widths):
            raise ValueError("Độ rộng cột phải là các số dương")
        if sum(widths) != int(table_width_dxa):
            raise ValueError("Tổng độ rộng cột phải bằng độ rộng bảng")

        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table_properties = table._tbl.tblPr
        cls._set_width(table_properties, "w:tblW", table_width_dxa)
        table_indent = cls._ensure_child(table_properties, "w:tblInd")
        table_indent.set(qn("w:type"), "dxa")
        table_indent.set(qn("w:w"), str(int(indent_dxa)))
        layout = cls._ensure_child(table_properties, "w:tblLayout")
        layout.set(qn("w:type"), "fixed")

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            grid_column = OxmlElement("w:gridCol")
            grid_column.set(qn("w:w"), str(width))
            grid.append(grid_column)

        for column_index, width in enumerate(widths):
            table.columns[column_index].width = Twips(width)
        for row in table.rows:
            if len(row.cells) != len(widths):
                raise ValueError("Fallback geometry không hỗ trợ hàng đã merge")
            for column_index, cell in enumerate(row.cells):
                width = widths[column_index]
                cell.width = Twips(width)
                cell_properties = cell._tc.get_or_add_tcPr()
                cls._set_width(cell_properties, "w:tcW", width)
                margins = cls._ensure_child(cell_properties, "w:tcMar")
                for side in ("top", "bottom", "start", "end"):
                    margin = cls._ensure_child(margins, f"w:{side}")
                    margin.set(qn("w:w"), str(int(cell_margins_dxa[side])))
                    margin.set(qn("w:type"), "dxa")


def table_helper_candidates() -> Iterable[Path]:
    """Ưu tiên override/repo, sau đó tìm bản Documents skill không khóa version."""

    configured = os.getenv("DE_GENESIS_TABLE_GEOMETRY_HELPER")
    if configured:
        yield Path(configured).expanduser()
    yield PROJECT_ROOT / "scripts" / "table_geometry.py"
    cache_root = Path.home() / ".codex" / "plugins" / "cache"
    if cache_root.is_dir():
        yield from sorted(
            cache_root.glob(
                "*/documents/*/skills/documents/scripts/table_geometry.py"
            ),
            reverse=True,
        )


def load_table_helper():
    for helper_path in table_helper_candidates():
        if not helper_path.is_file():
            continue
        spec = importlib.util.spec_from_file_location(
            "de_genesis_table_geometry", helper_path
        )
        if spec is None or spec.loader is None:
            continue
        module = importlib.util.module_from_spec(spec)
        try:
            spec.loader.exec_module(module)
        except (ImportError, OSError):
            continue
        if all(
            hasattr(module, name)
            for name in ("column_widths_from_weights", "apply_table_geometry")
        ):
            return module
    return FallbackTableGeometry


TABLE_GEOMETRY = load_table_helper()


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "C8D0DA", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def configure_document(document: Document) -> None:
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in style_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.keep_with_next = True

    code_style = document.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.12)

    caption_style = document.styles["Caption"]
    caption_style.font.name = "Calibri"
    caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption_style.font.size = Pt(9)
    caption_style.font.italic = True
    caption_style.font.color.rgb = MUTED
    caption_style.paragraph_format.space_before = Pt(3)
    caption_style.paragraph_format.space_after = Pt(8)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def populate_header(header) -> None:
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.paragraph_format.space_after = Pt(0)
        run = header_p.add_run("DE GENESIS  |  BÁO CÁO TUẦN 2")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        p_pr = header_p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "D7DBE2")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    populate_header(section.header)
    populate_header(section.even_page_header)
    add_page_number(section.footer.paragraphs[0])
    add_page_number(section.even_page_footer.paragraphs[0])


def configure_numbering(document: Document) -> tuple[int, int, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    next_abstract = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def add_definition(abstract_id: int, num_id: int, fmt: str, text_value: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(tabs)
        p_pr.append(ind)
        p_pr.append(spacing)
        lvl.append(start)
        lvl.append(num_fmt)
        lvl.append(lvl_text)
        lvl.append(suff)
        lvl.append(p_pr)
        abstract.append(lvl)
        first_num_index = next(
            (
                index
                for index, child in enumerate(numbering)
                if child.tag == qn("w:num")
            ),
            len(numbering),
        )
        numbering.insert(first_num_index, abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    bullet_num = next_num
    decimal_num = next_num + 1
    add_definition(next_abstract, bullet_num, "bullet", "•")
    add_definition(next_abstract + 1, decimal_num, "decimal", "%1.")
    return bullet_num, decimal_num, next_abstract + 1


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)


def create_numbering_instance(document: Document, abstract_num_id: int) -> int:
    numbering = document.part.numbering_part.element
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_num, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline_text(paragraph, text_value: str, *, default_size: float = 11) -> None:
    for part in INLINE_PATTERN.split(text_value):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(
                run,
                name="Consolas",
                size=max(default_size - 1, 8),
                color=DARK_BLUE,
            )
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=default_size)


def add_body_paragraph(document: Document, text_value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.widow_control = True
    add_inline_text(paragraph, text_value)


def add_list_item(document: Document, text_value: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    apply_numbering(paragraph, num_id)
    add_inline_text(paragraph, text_value)


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), VERY_LIGHT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "3")
        border.set(qn("w:color"), "D7DBE2")
        border.set(qn("w:space"), "4")
        borders.append(border)
    p_pr.append(borders)
    run = paragraph.add_run(code.rstrip())
    set_run_font(run, name="Consolas", size=8.5)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw_cells = lines[index].strip().strip("|").split("|")
        rows.append([cell.strip() for cell in raw_cells])
        index += 1
    if len(rows) >= 2 and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]
    ):
        rows.pop(1)
    return rows, index


def table_column_widths(rows: list[list[str]]) -> list[int]:
    column_count = len(rows[0])
    lengths = []
    for col in range(column_count):
        max_length = max(len(re.sub(r"[`*]", "", row[col])) for row in rows)
        lengths.append(max(8, min(max_length, 42)))
    weights = [length**0.65 for length in lengths]
    if column_count == 2 and lengths[0] <= 18 and lengths[1] >= 30:
        weights = [1.6, 4.9]
    return TABLE_GEOMETRY.column_widths_from_weights(weights, CONTENT_WIDTH_DXA)


def is_numeric_text(value: str) -> bool:
    cleaned = value.replace(".", "").replace(",", "").replace("%", "").strip()
    return bool(cleaned) and cleaned.replace("-", "").isdigit()


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        keep_table_row_together(row)
        if row_index == 0:
            repeat_table_header(row)
        for col_index, value in enumerate(row_values):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_shading(cell, LIGHT_GRAY if row_index == 0 else "FFFFFF")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT
                if row_index > 0 and is_numeric_text(value)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline_text(paragraph, value, default_size=9.2)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE

    TABLE_GEOMETRY.apply_table_geometry(
        table,
        table_column_widths(rows),
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS_DXA,
    )
    spacing = document.add_paragraph()
    spacing.paragraph_format.space_before = Pt(0)
    spacing.paragraph_format.space_after = Pt(2)


def get_font(size: int, bold: bool = False):
    configured = os.getenv(
        "DE_GENESIS_REPORT_FONT_BOLD" if bold else "DE_GENESIS_REPORT_FONT_REGULAR"
    )
    candidates = [Path(configured).expanduser()] if configured else []
    windows_root = os.getenv("WINDIR")
    if windows_root:
        candidates.append(
            Path(windows_root) / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf")
        )
    candidates.extend(
        [
            Path(
                "/usr/share/fonts/truetype/dejavu/"
                + ("DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf")
            ),
            Path(
                "/usr/share/fonts/truetype/liberation2/"
                + ("LiberationSans-Bold.ttf" if bold else "LiberationSans-Regular.ttf")
            ),
            Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf"),
        ]
    )
    for path in candidates:
        if path.is_file():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def rounded_box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    title: str,
    subtitle: str = "",
    *,
    fill: str = "#E8EEF5",
    outline: str = "#2E74B5",
    title_color: str = "#1F4D78",
) -> None:
    draw.rounded_rectangle(rect, radius=18, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = rect
    title_font = get_font(28, bold=True)
    body_font = get_font(20)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    title_width = title_box[2] - title_box[0]
    draw.text(
        ((x1 + x2 - title_width) / 2, y1 + 24),
        title,
        font=title_font,
        fill=title_color,
    )
    if subtitle:
        lines = subtitle.split("\n")
        y = y1 + 70
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=body_font)
            width = bbox[2] - bbox[0]
            draw.text(
                ((x1 + x2 - width) / 2, y),
                line,
                font=body_font,
                fill="#4B5563",
            )
            y += 28


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    color: str = "#6B7280",
    width: int = 5,
) -> None:
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 16 * direction, y2 - 10), (x2 - 16 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 10, y2 - 16 * direction), (x2 + 10, y2 - 16 * direction)]
    draw.polygon(points, fill=color)


def generate_pipeline_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 430), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        (40, 130, 330, 300, "9 CSV Olist", "Dữ liệu nguồn"),
        (390, 130, 680, 300, "olist_practice", "Làm sạch\n9 bảng"),
        (740, 130, 1030, 300, "olist_oltp", "3NF\n12 bảng"),
        (1090, 130, 1380, 300, "olist_olap", "Kimball\n7 dim + 4 fact"),
        (1440, 130, 1760, 300, "SQL / BI", "Báo cáo\nDashboard"),
    ]
    for index, (x1, y1, x2, y2, title, subtitle) in enumerate(boxes):
        rounded_box(draw, (x1, y1, x2, y2), title, subtitle)
        if index < len(boxes) - 1:
            draw_arrow(draw, (x2 + 5, (y1 + y2) // 2), (boxes[index + 1][0] - 8, (y1 + y2) // 2))
    image.save(path)


def generate_oltp_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(image)
    positions = {
        "customers": (40, 90, 380, 220),
        "customer_addresses": (500, 90, 900, 220),
        "postal_locations": (1050, 90, 1450, 220),
        "orders": (500, 390, 900, 530),
        "order_statuses": (40, 390, 380, 530),
        "order_items": (500, 770, 900, 910),
        "products": (1050, 640, 1450, 770),
        "product_categories": (1450, 820, 1770, 950),
        "sellers": (1050, 900, 1450, 1030),
        "order_payments": (40, 650, 380, 790),
        "payment_methods": (40, 900, 380, 1030),
        "order_reviews": (1450, 350, 1770, 490),
    }
    connections = [
        ("customers", "customer_addresses"),
        ("postal_locations", "customer_addresses"),
        ("customer_addresses", "orders"),
        ("order_statuses", "orders"),
        ("orders", "order_items"),
        ("products", "order_items"),
        ("sellers", "order_items"),
        ("product_categories", "products"),
        ("orders", "order_payments"),
        ("payment_methods", "order_payments"),
        ("orders", "order_reviews"),
        ("postal_locations", "sellers"),
    ]
    for source, target in connections:
        s = positions[source]
        t = positions[target]
        start = ((s[0] + s[2]) // 2, (s[1] + s[3]) // 2)
        end = ((t[0] + t[2]) // 2, (t[1] + t[3]) // 2)
        draw_arrow(draw, start, end, color="#B1BAC6", width=4)
    for name, rect in positions.items():
        fill = "#E8EEF5" if name in {"orders", "order_items", "order_payments", "order_reviews"} else "#F2F4F7"
        rounded_box(draw, rect, name, fill=fill)
    image.save(path)


def generate_star_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1120), "white")
    draw = ImageDraw.Draw(image)
    facts = {
        "fact_sales": (700, 100, 1100, 240),
        "fact_payments": (700, 360, 1100, 500),
        "fact_order_lifecycle": (700, 620, 1100, 760),
        "fact_reviews": (700, 880, 1100, 1020),
    }
    dims = {
        "dim_date": (40, 80, 410, 210),
        "dim_customer": (40, 300, 410, 430),
        "dim_location": (40, 520, 410, 650),
        "dim_order_status": (40, 740, 410, 870),
        "dim_product": (1390, 120, 1760, 250),
        "dim_seller": (1390, 330, 1760, 460),
        "dim_payment_method": (1390, 560, 1760, 690),
    }
    shared_dims = ["dim_date", "dim_customer", "dim_location", "dim_order_status"]
    for dim in shared_dims:
        d = dims[dim]
        start = ((d[0] + d[2]) // 2, (d[1] + d[3]) // 2)
        for fact_rect in facts.values():
            end = (fact_rect[0], (fact_rect[1] + fact_rect[3]) // 2)
            draw.line([start, end], fill="#C1C9D3", width=3)
    for dim, fact in (
        ("dim_product", "fact_sales"),
        ("dim_seller", "fact_sales"),
        ("dim_payment_method", "fact_payments"),
    ):
        d = dims[dim]
        f = facts[fact]
        draw.line(
            [((d[0] + d[2]) // 2, (d[1] + d[3]) // 2), (f[2], (f[1] + f[3]) // 2)],
            fill="#C1C9D3",
            width=3,
        )
    for name, rect in dims.items():
        rounded_box(draw, rect, name, "Conformed dimension", fill="#F2F4F7")
    for name, rect in facts.items():
        rounded_box(
            draw,
            rect,
            name,
            "Fact table",
            fill="#DCEAF7",
            outline="#1F4D78",
        )
    image.save(path)


def generate_diagrams() -> list[Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = [
        ASSET_DIR / "pipeline.png",
        ASSET_DIR / "oltp_erd.png",
        ASSET_DIR / "star_schema.png",
    ]
    generate_pipeline_diagram(paths[0])
    generate_oltp_diagram(paths[1])
    generate_star_diagram(paths[2])
    return paths


def set_picture_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.2))
    set_picture_alt_text(shape, caption)
    caption_p = document.add_paragraph(caption, style="Caption")
    caption_p.paragraph_format.keep_with_next = False


def add_cover(document: Document) -> None:
    for _ in range(4):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("DE GENESIS  |  DATA ENGINEERING")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("BÁO CÁO TUẦN 2")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run(
        "THIẾT KẾ CƠ SỞ DỮ LIỆU OLTP\n"
        "VÀ KHO DỮ LIỆU OLAP THEO KIMBALL"
    )
    set_run_font(run, size=16, color=DARK_BLUE, bold=True)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(28)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    metadata = [
        ("Dự án", "DE Genesis - Lộ trình thực hành Data Engineering"),
        ("Bộ dữ liệu", "Brazilian E-Commerce Public Dataset by Olist"),
        ("Công nghệ", "Python, PostgreSQL 16, SQLAlchemy, Docker Compose"),
        ("Thời gian", "Tuần 2 - hoàn thành ngày 04/07/2026"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=MUTED, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)

    document.add_page_break()


def add_heading(document: Document, level: int, title: str) -> None:
    title = title.replace("`", "").replace("**", "")
    style_name = f"Heading {min(level, 3)}"
    paragraph = document.add_paragraph(style=style_name)
    paragraph.paragraph_format.widow_control = True
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    run = paragraph.add_run(title)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[min(level, 3)],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[min(level, 3)],
        bold=True,
    )


def convert_markdown_to_docx(document: Document, markdown: str, diagram_paths: list[Path]) -> None:
    lines = markdown.splitlines()
    start = 0
    for index, line in enumerate(lines):
        if line.strip() == "---":
            start = index + 1
            break
    lines = lines[start:]

    bullet_num, decimal_num, decimal_abstract_id = configure_numbering(document)
    diagram_captions = [
        "Hình 1. Kiến trúc dữ liệu tổng thể của tuần 2",
        "Hình 2. Quan hệ chính trong mô hình OLTP",
        "Hình 3. Bốn star schema dùng chung conformed dimensions",
    ]
    diagram_index = 0
    index = 0
    paragraph_buffer: list[str] = []
    toc_seen = False
    appendix_started = False

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_body_paragraph(document, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if language == "mermaid":
                add_figure(
                    document,
                    diagram_paths[diagram_index],
                    diagram_captions[diagram_index],
                )
                diagram_index += 1
            else:
                add_code_block(document, "\n".join(code_lines))
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_rows, next_index = parse_markdown_table(lines, index)
            add_table(document, table_rows)
            index = next_index
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            if title == "MỤC LỤC":
                document.add_page_break()
                toc_seen = True
            if title == "6.3.1. Conformed dimensions và Bus Architecture":
                document.add_page_break()
            if level == 1 and re.match(r"\d+\.", title) and toc_seen:
                document.add_page_break()
                toc_seen = False
            elif level == 1 and title.startswith("PHỤ LỤC"):
                if appendix_started:
                    document.add_page_break()
                appendix_started = True
            add_heading(document, level, title)
            index += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            flush_paragraph()
            add_list_item(document, re.sub(r"^[-*]\s+", "", stripped), bullet_num)
            index += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            flush_paragraph()
            list_num_id = create_numbering_instance(document, decimal_abstract_id)
            while index < len(lines):
                current = lines[index].strip()
                current_match = re.match(r"^\d+\.\s+(.+)$", current)
                if current_match:
                    add_list_item(document, current_match.group(1), list_num_id)
                    index += 1
                    continue
                if (
                    not current
                    and index + 1 < len(lines)
                    and re.match(r"^\d+\.\s+", lines[index + 1].strip())
                ):
                    index += 1
                    continue
                break
            continue

        if stripped == "---":
            flush_paragraph()
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(stripped.rstrip("  "))
        index += 1

    flush_paragraph()


def add_core_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = "Báo cáo tuần 2 - Thiết kế OLTP và OLAP theo Kimball"
    properties.subject = "Dự án DE Genesis"
    properties.author = "DE Genesis"
    properties.keywords = "Data Engineering, OLTP, OLAP, Kimball, PostgreSQL, Olist"
    properties.comments = "Báo cáo được tạo từ kết quả chạy và đối soát thực tế."


def build() -> Path:
    diagram_paths = generate_diagrams()
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    add_core_properties(document)
    add_cover(document)
    convert_markdown_to_docx(document, markdown, diagram_paths)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
