from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
MARKDOWN_PATH = HERE / "bao_cao_tuan_6.md"
OUTPUT_PATH = HERE / "bao_cao_tuan_6.docx"
ASSET_DIR = HERE.parents[2] / "output" / "week6" / "report_assets"
ARCHITECTURE_PATH = ASSET_DIR / "week6_architecture.png"

# Preset: standard_business_brief.
# Header template: editorial_cover.
CONTENT_DXA = 9360
NAVY = RGBColor(25, 55, 90)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
HEADER_FILL = "F2F4F7"


def set_font(run, size=11, color=BLACK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def configure(document: Document) -> tuple[int, int]:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    headings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in headings.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0, line=1.0)
    set_font(header.add_run("DE GENESIS  |  BÁO CÁO TUẦN 6"), 8.5, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0, line=1.0)
    set_font(footer.add_run("Trang "), 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    return create_numbering(document, bullet=True), create_numbering(document, bullet=False)


def create_numbering(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([level, num])
    p_pr.append(num_pr)
    set_paragraph_spacing(paragraph, after=8, line=1.167)


def add_cover(document: Document) -> None:
    for _ in range(5):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(14)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, after=18, line=1.0)
    set_font(kicker.add_run("DE GENESIS  |  DATA ENGINEERING LAB"), 11, BLUE, True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=10, line=1.0)
    set_font(title.add_run("BÁO CÁO TUẦN 6"), 30, NAVY, True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=34, line=1.1)
    set_font(
        subtitle.add_run("PRODUCTION PIPELINE,\nMONITORING VÀ CẢNH BÁO"),
        15.5,
        DARK_BLUE,
        True,
    )
    metadata = [
        ("Dự án", "DE Genesis - Lộ trình thực hành Data Engineering"),
        ("Phạm vi", "Production pipeline và observability đầu cuối"),
        ("Công nghệ", "Airflow 2.9.3, Spark 3.5.1, PostgreSQL 16, Prometheus 2.52, Grafana 10.4"),
        ("Ngày báo cáo", "31/07/2026"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=4, line=1.0)
        set_font(paragraph.add_run(f"{label}: "), 10.5, MUTED, True)
        set_font(paragraph.add_run(value), 10.5)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(46)
    note.paragraph_format.space_after = Pt(0)
    set_font(
        note.add_run("Báo cáo dựa trên mã nguồn, audit database và nghiệm thu Docker thực tế."),
        9.5,
        MUTED,
        italic=True,
    )
    document.add_page_break()


def load_image_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_box(draw, rect, title, subtitle, fill):
    draw.rounded_rectangle(rect, radius=18, fill=fill, outline="#2E74B5", width=3)
    x1, y1, x2, _ = rect
    title_font = load_image_font(24, True)
    subtitle_font = load_image_font(18)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((x1 + x2 - title_box[2]) / 2, y1 + 25), title, font=title_font, fill="#19375A")
    for index, line in enumerate(subtitle.split("\n")):
        line_box = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(
            ((x1 + x2 - line_box[2]) / 2, y1 + 72 + index * 25),
            line,
            font=subtitle_font,
            fill="#5A636E",
        )


def draw_arrow(draw, start, end):
    draw.line([start, end], fill="#2E74B5", width=5)
    draw.polygon(
        [(end[0], end[1]), (end[0] - 15, end[1] - 9), (end[0] - 15, end[1] + 9)],
        fill="#2E74B5",
    )


def generate_architecture() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    title = "KIẾN TRÚC PRODUCTION PIPELINE VÀ OBSERVABILITY TUẦN 6"
    title_box = draw.textbbox((0, 0), title, font=load_image_font(36, True))
    draw.text(
        ((1800 - title_box[2]) / 2, 25),
        title,
        font=load_image_font(36, True),
        fill="#19375A",
    )
    top = [
        (35, 170, 260, 345, "Promotion API", "Incremental\nwindow"),
        (330, 170, 555, 345, "Airflow", "Audit + retry\nbackfill"),
        (625, 170, 850, 345, "Raw + DQ", "Validate\nblocking gate"),
        (920, 170, 1145, 345, "Spark", "Snapshot\nrefresh"),
        (1215, 170, 1440, 345, "Curated + DQ", "Grain +\nfinance rules"),
        (1510, 170, 1765, 345, "Watermark", "Chỉ cập nhật\nsau success"),
    ]
    for index, (*rect, box_title, subtitle) in enumerate(top):
        draw_box(draw, tuple(rect), box_title, subtitle, "#F2F4F7")
        if index < len(top) - 1:
            draw_arrow(
                draw,
                (rect[2] + 10, (rect[1] + rect[3]) // 2),
                (top[index + 1][0] - 10, (rect[1] + rect[3]) // 2),
            )
    bottom = [
        (330, 590, 650, 770, "Metrics exporter", "Audit database\nDependency health"),
        (740, 590, 1060, 770, "Prometheus", "4 scrape target\n6 alert rule"),
        (1150, 590, 1470, 770, "Grafana", "Provisioned data source\n5 dashboard panel"),
    ]
    for index, (*rect, box_title, subtitle) in enumerate(bottom):
        draw_box(draw, tuple(rect), box_title, subtitle, "#E8EEF5")
        if index < len(bottom) - 1:
            draw_arrow(
                draw,
                (rect[2] + 10, (rect[1] + rect[3]) // 2),
                (bottom[index + 1][0] - 10, (rect[1] + rect[3]) // 2),
            )
    draw.line([(1635, 345), (1635, 500), (490, 500), (490, 590)], fill="#2E74B5", width=5)
    draw.polygon([(490, 590), (480, 574), (500, 574)], fill="#2E74B5")
    image.save(ARCHITECTURE_PATH)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            set_font(paragraph.add_run(part[1:-1]), 9.5, DARK_BLUE, name="Consolas")
        elif part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), 11, bold=True)
        else:
            set_font(paragraph.add_run(part), 11)


def table_widths(columns: int) -> list[int]:
    patterns = {
        2: [2700, 6660],
        3: [1900, 3100, 4360],
        4: [1800, 2520, 2520, 2520],
    }
    if columns in patterns:
        return patterns[columns]
    widths = [CONTENT_DXA // columns] * columns
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(document: Document, rows):
    columns = len(rows[0])
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.05)
            set_font(
                paragraph.add_run(value),
                9.5,
                WHITE if row_index == 0 else BLACK,
                row_index == 0,
            )
            if row_index == 0:
                shade_cell(cell, "2E74B5")
    set_table_geometry(table, table_widths(columns))
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def add_architecture(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    picture_run = paragraph.add_run()
    picture_run.add_picture(str(ARCHITECTURE_PATH), width=Inches(6.35))
    doc_property = picture_run._r.xpath(".//wp:docPr")[0]
    doc_property.set(
        "descr",
        "Sơ đồ Production Pipeline từ Promotion API qua Airflow, raw quality gate, Spark, curated quality gate và watermark; metrics được đưa vào Prometheus và Grafana.",
    )
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    set_font(
        caption.add_run("Hình 1. Kiến trúc production pipeline và observability tuần 6"),
        9,
        MUTED,
        italic=True,
    )


def parse_body(document: Document, markdown: str, bullet_num_id: int, decimal_num_id: int) -> None:
    lines = markdown.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("# 1."))
    index = start
    major_count = 0
    paragraph_buffer: list[str] = []

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index].rstrip()
        if line == "[[ARCHITECTURE]]":
            flush()
            add_architecture(document)
        elif line.startswith("# "):
            flush()
            major_count += 1
            if major_count > 1 and major_count in {4, 7, 10, 12, 14, 16}:
                document.add_page_break()
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            flush()
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            flush()
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            flush()
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, bullet_num_id)
            inline(paragraph, line[2:])
        elif re.match(r"^\d+\. ", line):
            flush()
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, decimal_num_id)
            inline(paragraph, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("|"):
            flush()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|\s*:?-+", row)
            ]
            add_table(document, parsed)
            continue
        elif not line:
            flush()
        else:
            paragraph_buffer.append(line)
        index += 1
    flush()


def audit_document(path: Path) -> None:
    document = Document(path)
    section = document.sections[0]
    expected = {
        "page_width": Inches(8.5),
        "page_height": Inches(11),
        "top_margin": Inches(1),
        "bottom_margin": Inches(1),
        "left_margin": Inches(1),
        "right_margin": Inches(1),
    }
    for attribute, value in expected.items():
        actual = getattr(section, attribute)
        if abs(actual - value) > 5:
            raise AssertionError(f"Section token sai: {attribute}={actual}")
    if not document.inline_shapes:
        raise AssertionError("Báo cáo thiếu sơ đồ kiến trúc")
    if not document.tables:
        raise AssertionError("Báo cáo thiếu bảng")
    for table in document.tables:
        grid_widths = [
            int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        if sum(grid_widths) != CONTENT_DXA:
            raise AssertionError(f"Table grid không đủ {CONTENT_DXA} DXA: {grid_widths}")
        for row in table.rows:
            cell_widths = [
                int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w")))
                for cell in row.cells
            ]
            if cell_widths != grid_widths:
                raise AssertionError("tcW không khớp tblGrid")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for forbidden in ("[[ARCHITECTURE]]", "TODO", "PLACEHOLDER"):
        if forbidden in text:
            raise AssertionError(f"Còn token nội bộ: {forbidden}")


def build() -> None:
    generate_architecture()
    document = Document()
    bullet_num_id, decimal_num_id = configure(document)
    add_cover(document)
    parse_body(
        document,
        MARKDOWN_PATH.read_text(encoding="utf-8"),
        bullet_num_id,
        decimal_num_id,
    )
    document.core_properties.title = "Báo cáo thực hành tuần 6 - DE Genesis"
    document.core_properties.subject = "Production pipeline, monitoring và cảnh báo"
    document.core_properties.author = "DE Genesis"
    document.core_properties.keywords = "Airflow, Spark, Prometheus, Grafana, Data Quality"
    document.save(OUTPUT_PATH)
    audit_document(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
