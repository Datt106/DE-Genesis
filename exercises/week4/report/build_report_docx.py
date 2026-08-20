from __future__ import annotations

import importlib.util
import re
from pathlib import Path
from typing import Callable

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
PROJECT_ROOT = HERE.parents[2]
MARKDOWN_PATH = HERE / "bao_cao_tuan_3_4.md"
OUTPUT_PATH = HERE / "bao_cao_tuan_3_4.docx"
ASSET_DIR = PROJECT_ROOT / "output" / "week4" / "report_assets"
BASE_BUILDER_PATH = HERE.parents[1] / "week2" / "report" / "build_report_docx.py"

NAVY = RGBColor(25, 55, 90)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
BLACK = RGBColor(0, 0, 0)


def load_base_builder():
    spec = importlib.util.spec_from_file_location("week2_report_builder", BASE_BUILDER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Không thể nạp builder nền: {BASE_BUILDER_PATH}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


BASE = load_base_builder()


def configure_document(document: Document) -> None:
    """Áp dụng preset standard_business_brief và header riêng của báo cáo."""

    BASE.configure_document(document)
    document.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    section = document.sections[0]
    for header in (section.header, section.even_page_header):
        paragraph = header.paragraphs[0]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        p_pr = paragraph._p.get_or_add_pPr()
        border = p_pr.find(qn("w:pBdr"))
        if border is not None:
            p_pr.remove(border)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("DE GENESIS  |  BÁO CÁO TUẦN 3 & 4")
        BASE.set_run_font(run, size=8.5, color=MUTED, bold=True)


def add_cover(document: Document) -> None:
    """Editorial cover: title stack giữa trang, nhiều khoảng thở, không metadata grid."""

    for _ in range(5):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(14)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("DE GENESIS  |  DATA ENGINEERING LAB")
    BASE.set_run_font(run, size=11, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("BÁO CÁO TUẦN 3 & 4")
    BASE.set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run(
        "SPARK BATCH, HDFS, PARQUET/ORC\n"
        "KAFKA VÀ SPARK STRUCTURED STREAMING"
    )
    BASE.set_run_font(run, size=15.5, color=DARK_BLUE, bold=True)

    metadata = [
        ("Dự án", "DE Genesis - Lộ trình thực hành Data Engineering"),
        ("Phạm vi", "Tuần 3 và Tuần 4"),
        ("Môi trường", "Docker Compose, Spark 3.5.1, Hadoop 3.2.1, Kafka 7.6.1"),
        ("Ngày hoàn thành", "12/07/2026"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        label_run = paragraph.add_run(f"{label}: ")
        BASE.set_run_font(label_run, size=10.5, color=MUTED, bold=True)
        value_run = paragraph.add_run(value)
        BASE.set_run_font(value_run, size=10.5, color=BLACK)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(54)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Báo cáo được lập từ mã nguồn và kết quả chạy thực tế trong repo.")
    BASE.set_run_font(run, size=9.5, color=MUTED, italic=True)

    document.add_page_break()


def centered_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    *,
    font,
    fill: str,
) -> None:
    box = draw.textbbox((0, 0), text, font=font)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1]), text, font=font, fill=fill)


def title_on_image(draw: ImageDraw.ImageDraw, text: str) -> None:
    centered_text(
        draw,
        (900, 24),
        text,
        font=BASE.get_font(34, bold=True),
        fill="#19375A",
    )


def label_lane(draw: ImageDraw.ImageDraw, y: int, text: str, fill: str) -> None:
    draw.rounded_rectangle((35, y, 220, y + 66), radius=16, fill=fill)
    centered_text(
        draw,
        (127, y + 17),
        text,
        font=BASE.get_font(23, bold=True),
        fill="#FFFFFF",
    )


def generate_overview(path: Path) -> None:
    image = Image.new("RGB", (1800, 780), "white")
    draw = ImageDraw.Draw(image)
    title_on_image(draw, "KIẾN TRÚC THỰC HÀNH TUẦN 3 VÀ TUẦN 4")

    label_lane(draw, 140, "TUẦN 3", "#2E74B5")
    top_boxes = [
        (260, 112, 520, 292, "sales.csv", "Local / HDFS"),
        (580, 112, 840, 292, "Spark batch", "Kiểm tra\nBiến đổi"),
        (900, 112, 1160, 292, "Aggregate", "Category\nRegion / Daily"),
        (1220, 112, 1480, 292, "Parquet / ORC", "Curated\nPartition"),
        (1540, 112, 1765, 292, "Quality", "JSON\nĐối soát"),
    ]
    for index, (*rect, title, subtitle) in enumerate(top_boxes):
        BASE.rounded_box(draw, tuple(rect), title, subtitle)
        if index < len(top_boxes) - 1:
            BASE.draw_arrow(draw, (rect[2] + 5, 202), (top_boxes[index + 1][0] - 8, 202))

    label_lane(draw, 500, "TUẦN 4", "#1F4D78")
    bottom_boxes = [
        (260, 472, 520, 652, "Producer", "JSON events\nacks=all"),
        (580, 472, 840, 652, "Kafka", "Topic\nPartition / Offset"),
        (900, 472, 1160, 652, "Streaming", "Parse / Filter\nEvent time"),
        (1220, 472, 1480, 652, "Window", "Watermark\nAggregate"),
        (1540, 472, 1765, 652, "Parquet", "Output\nCheckpoint"),
    ]
    for index, (*rect, title, subtitle) in enumerate(bottom_boxes):
        BASE.rounded_box(draw, tuple(rect), title, subtitle, fill="#F2F4F7")
        if index < len(bottom_boxes) - 1:
            BASE.draw_arrow(draw, (rect[2] + 5, 562), (bottom_boxes[index + 1][0] - 8, 562))

    centered_text(
        draw,
        (900, 710),
        "Hạ tầng dùng chung: Docker Compose • Spark master/worker • /workspace bind mount",
        font=BASE.get_font(22),
        fill="#5A636E",
    )
    image.save(path)


def generate_flow(
    path: Path,
    heading: str,
    steps: list[tuple[str, str]],
    *,
    fill: str,
) -> None:
    image = Image.new("RGB", (1900, 540), "white")
    draw = ImageDraw.Draw(image)
    centered_text(
        draw,
        (950, 25),
        heading,
        font=BASE.get_font(34, bold=True),
        fill="#19375A",
    )
    gap = 22
    box_width = 280
    start_x = 42
    y1, y2 = 150, 385
    for index, (title, subtitle) in enumerate(steps):
        x1 = start_x + index * (box_width + gap)
        x2 = x1 + box_width
        BASE.rounded_box(
            draw,
            (x1, y1, x2, y2),
            title,
            subtitle,
            fill=fill,
        )
        if index < len(steps) - 1:
            BASE.draw_arrow(draw, (x2 + 4, 267), (x2 + gap - 5, 267))
    image.save(path)


def generate_week3_revenue(path: Path) -> None:
    values = [("Electronics", 2980), ("Furniture", 655), ("Office", 49)]
    image = Image.new("RGB", (1800, 720), "white")
    draw = ImageDraw.Draw(image)
    title_on_image(draw, "DOANH THU THEO DANH MỤC - TUẦN 3")
    left, right = 330, 1650
    top, bar_height, gap = 150, 105, 70
    maximum = max(value for _, value in values)
    label_font = BASE.get_font(28, bold=True)
    value_font = BASE.get_font(26, bold=True)
    colors = ["#2E74B5", "#6D9ECF", "#B9D3EA"]
    for index, ((label, value), color) in enumerate(zip(values, colors)):
        y = top + index * (bar_height + gap)
        draw.text((55, y + 30), label, font=label_font, fill="#1F4D78")
        width = max(18, int((right - left) * value / maximum))
        draw.rounded_rectangle((left, y, left + width, y + bar_height), radius=18, fill=color)
        formatted = f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        draw.text((left + width + 20, y + 32), formatted, font=value_font, fill="#19375A")
    centered_text(
        draw,
        (900, 650),
        "Tổng doanh thu: 3.684,00 • Electronics chiếm 80,89%",
        font=BASE.get_font(24),
        fill="#5A636E",
    )
    image.save(path)


def generate_week4_results(path: Path) -> None:
    values = [("catalog", 11, 2), ("checkout", 5, 2), ("payment", 4, 0)]
    image = Image.new("RGB", (1800, 720), "white")
    draw = ImageDraw.Draw(image)
    title_on_image(draw, "REQUEST VÀ SERVER ERROR THEO SERVICE - TUẦN 4")
    left, scale = 410, 92
    top, group_gap = 155, 165
    label_font = BASE.get_font(28, bold=True)
    value_font = BASE.get_font(24, bold=True)
    for index, (service, requests, errors) in enumerate(values):
        y = top + index * group_gap
        draw.text((70, y + 22), service, font=label_font, fill="#1F4D78")
        draw.rounded_rectangle(
            (left, y, left + requests * scale, y + 62),
            radius=14,
            fill="#2E74B5",
        )
        draw.text(
            (left + requests * scale + 16, y + 17),
            f"{requests} request",
            font=value_font,
            fill="#19375A",
        )
        if errors:
            draw.rounded_rectangle(
                (left, y + 76, left + errors * scale, y + 122),
                radius=12,
                fill="#9B1C1C",
            )
            draw.text(
                (left + errors * scale + 16, y + 83),
                f"{errors} lỗi 5xx",
                font=value_font,
                fill="#9B1C1C",
            )
        else:
            draw.text((left, y + 83), "0 lỗi 5xx", font=value_font, fill="#5A636E")
    centered_text(
        draw,
        (900, 660),
        "20 request • 18 nhóm cửa sổ • 4 server error • latency trung bình 546,85 ms",
        font=BASE.get_font(23),
        fill="#5A636E",
    )
    image.save(path)


def generate_diagrams() -> dict[str, tuple[Path, str]]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "overview": (
            ASSET_DIR / "overview.png",
            "Hình 1. Kiến trúc thực hành tuần 3 và tuần 4",
        ),
        "week3_flow": (
            ASSET_DIR / "week3_flow.png",
            "Hình 2. Luồng xử lý Spark batch tuần 3",
        ),
        "week3_revenue": (
            ASSET_DIR / "week3_revenue.png",
            "Hình 3. Doanh thu theo danh mục từ kết quả tuần 3",
        ),
        "week4_flow": (
            ASSET_DIR / "week4_flow.png",
            "Hình 4. Luồng Kafka và Spark Structured Streaming tuần 4",
        ),
        "week4_results": (
            ASSET_DIR / "week4_results.png",
            "Hình 5. Request và server error theo service trong lần kiểm chứng",
        ),
    }
    generate_overview(figures["overview"][0])
    generate_flow(
        figures["week3_flow"][0],
        "LUỒNG XỬ LÝ SPARK BATCH TUẦN 3",
        [
            ("CSV", "Header\nRaw schema"),
            ("Quality", "Dedupe\nValidation"),
            ("Curated", "Cast type\nline_amount"),
            ("Aggregate", "Category\nRegion / Daily"),
            ("Storage", "Parquet\nORC / HDFS"),
            ("Reconcile", "Read back\nQuality JSON"),
        ],
        fill="#E8EEF5",
    )
    generate_week3_revenue(figures["week3_revenue"][0])
    generate_flow(
        figures["week4_flow"][0],
        "LUỒNG XỬ LÝ STRUCTURED STREAMING TUẦN 4",
        [
            ("Producer", "Seed / Count\nacks=all"),
            ("Kafka", "Topic\nOffset"),
            ("Parse", "JSON schema\nFilter"),
            ("Event time", "Watermark\nWindow"),
            ("Aggregate", "Request\nLatency / 5xx"),
            ("Sink", "Parquet\nCheckpoint"),
        ],
        fill="#F2F4F7",
    )
    generate_week4_results(figures["week4_results"][0])
    return figures


def add_body_paragraph(document: Document, text_value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.widow_control = True
    BASE.add_inline_text(paragraph, text_value)


def convert_markdown_to_docx(
    document: Document,
    markdown: str,
    figures: dict[str, tuple[Path, str]],
) -> None:
    lines = markdown.splitlines()
    bullet_num, _decimal_num, decimal_abstract_id = BASE.configure_numbering(document)
    index = 0
    paragraph_buffer: list[str] = []
    forced_page_break_titles = {
        "TÓM TẮT ĐIỀU HÀNH",
        "PHỤ LỤC A - RUNBOOK CHẠY LẠI TOÀN BỘ BÀI THỰC HÀNH",
        "TÀI LIỆU THAM KHẢO",
    }

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_body_paragraph(
                document,
                " ".join(part.strip() for part in paragraph_buffer),
            )
            paragraph_buffer.clear()

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        figure_match = re.fullmatch(r"<!--\s*FIGURE:([a-z0-9_]+)\s*-->", stripped)
        if figure_match:
            flush_paragraph()
            key = figure_match.group(1)
            image_path, caption = figures[key]
            BASE.add_figure(document, image_path, caption)
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            BASE.add_code_block(document, "\n".join(code_lines))
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_rows, next_index = BASE.parse_markdown_table(lines, index)
            BASE.add_table(document, table_rows)
            index = next_index
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            if level == 1 and title in forced_page_break_titles:
                document.add_page_break()
            BASE.add_heading(document, level, title)
            index += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            flush_paragraph()
            BASE.add_list_item(
                document,
                re.sub(r"^[-*]\s+", "", stripped),
                bullet_num,
            )
            index += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            flush_paragraph()
            list_num_id = BASE.create_numbering_instance(document, decimal_abstract_id)
            while index < len(lines):
                current = lines[index].strip()
                current_match = re.match(r"^\d+\.\s+(.+)$", current)
                if current_match:
                    BASE.add_list_item(document, current_match.group(1), list_num_id)
                    index += 1
                    continue
                if (
                    not current
                    and index + 1 < len(lines)
                    and re.match(r"^\d+\.\s+", lines[index + 1].strip())
                ):
                    index += 1
                    continue
                break
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(stripped.rstrip("  "))
        index += 1

    flush_paragraph()


def add_core_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = "Báo cáo tuần 3 và 4 - Spark, HDFS, Kafka và Streaming"
    properties.subject = "Dự án DE Genesis"
    properties.author = "DE Genesis"
    properties.keywords = (
        "Data Engineering, Spark, HDFS, Parquet, ORC, Kafka, Structured Streaming"
    )
    properties.comments = "Báo cáo được tạo từ mã nguồn và kết quả chạy thực tế."


def build() -> Path:
    figures = generate_diagrams()
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    add_core_properties(document)
    add_cover(document)
    convert_markdown_to_docx(document, markdown, figures)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
