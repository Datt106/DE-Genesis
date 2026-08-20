from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
MARKDOWN_PATH = HERE / "bao_cao_tuan_1.md"
OUTPUT_PATH = HERE / "bao_cao_tuan_1.docx"

NAVY = RGBColor(25, 55, 90)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)


def set_font(run, *, size=10.5, color=BLACK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=110, bottom=80, end=110) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(paragraph.add_run("DE Genesis  |  "), size=8.5, color=MUTED)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    set_font(run, size=8.5, color=MUTED)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, DARK_BLUE, 11, 5),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("DE GENESIS  |  BÁO CÁO TUẦN 1"), size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_cover(document: Document) -> None:
    for _ in range(5):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(14)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_font(kicker.add_run("DE GENESIS  |  DATA ENGINEERING ROADMAP"), size=10.5, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("BÁO CÁO TUẦN 1"), size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    set_font(
        subtitle.add_run("PYTHON, JAVA, SQL, LINUX,\nPOSTGRESQL VÀ MYSQL"),
        size=15,
        color=DARK_BLUE,
        bold=True,
    )

    for label, value in (
        ("Dự án", "DE Genesis - Lộ trình thực hành Data Engineering"),
        ("Dữ liệu", "Brazilian E-Commerce Public Dataset by Olist"),
        ("Môi trường", "Docker local, PostgreSQL 16 và MySQL 8.0"),
        ("Ngày đồng bộ", "14/08/2026"),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        set_font(paragraph.add_run(f"{label}: "), size=10.5, color=MUTED, bold=True)
        set_font(paragraph.add_run(value), size=10.5)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(42)
    set_font(
        note.add_run("Báo cáo đã đồng bộ với implementation và checklist roadmap."),
        size=9.5,
        color=MUTED,
        italic=True,
    )
    document.add_page_break()


def add_inline(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), size=9.5, color=DARK_BLUE, name="Consolas")
        else:
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]))


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    properties.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_font(run, size=8.5, color=NAVY, name="Consolas")


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    available = 6.7
    widths = [available / columns] * columns
    if columns == 2:
        widths = [available * 0.42, available * 0.58]
    elif columns == 3:
        widths = [available * 0.25, available * 0.375, available * 0.375]

    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value)
            if row_index == 0:
                set_cell_shading(cell, "2E74B5")
                for run in paragraph.runs:
                    run.font.color.rgb = WHITE
                    run.bold = True
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F2F4F7")
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 1  # Tiêu đề đã nằm trên trang bìa.
    paragraph_buffer: list[str] = []

    def flush() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            add_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            flush()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
        elif line.startswith("## "):
            flush()
            heading = line[3:]
            if heading.startswith("7."):
                document.add_page_break()
            document.add_heading(heading, level=1)
        elif line.startswith("### "):
            flush()
            document.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            flush()
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline(paragraph, line[2:])
        elif line.startswith("|"):
            flush()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|\s*:?-", row)
            ]
            add_table(document, rows)
            continue
        elif not line:
            flush()
        else:
            paragraph_buffer.append(line)
        index += 1
    flush()


def audit(path: Path) -> None:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    for required in ("PostgreSQL 16", "MySQL 8.0", "Stored procedure", "Trigger", "EXPLAIN"):
        if required not in text:
            raise AssertionError(f"Báo cáo thiếu nội dung: {required}")
    if len(document.tables) < 2:
        raise AssertionError("Báo cáo thiếu bảng đối soát")
    if any(token in text for token in ("```", "| ---", "TODO", "PLACEHOLDER")):
        raise AssertionError("Báo cáo còn token Markdown nội bộ")


def build() -> None:
    document = Document()
    configure(document)
    add_cover(document)
    parse_markdown(document, MARKDOWN_PATH.read_text(encoding="utf-8"))
    document.core_properties.title = "Báo cáo tuần 1 - DE Genesis"
    document.core_properties.subject = "Python, Java, SQL, Linux, PostgreSQL và MySQL"
    document.core_properties.author = "DE Genesis"
    document.core_properties.keywords = "Python, Java, SQL, Linux, PostgreSQL, MySQL, Olist"
    document.save(OUTPUT_PATH)
    audit(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
