import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PRIMARY_NAVY = "1F4E79"
SECONDARY_BLUE = "2F5597"
TEXT_DARK = "262626"
LIGHT_BG = "F2F4F7"
ZEBRA_BG = "F9FAFB"
BORDER_GRAY = "D9D9D9"
CODE_BG = "F4F4F4"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def add_callout(doc, text, title=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, LIGHT_BG)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{PRIMARY_NAVY}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        r_title = p.add_run(f"{title}\n")
        r_title.bold = True
        r_title.font.name = "Calibri"
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = RGBColor(31, 78, 121)
    
    parse_formatted_text(p, text, italic_default=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, CODE_BG)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY}"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY}"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(40, 40, 40)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def parse_formatted_text(paragraph, text, italic_default=False, font_size=10.5):
    # Regex split for bold (**text**), italic (*text*), code (`text`), markdown link ([text](url))
    # Tokens matching:
    # 1. Links: \[(.*?)\]\((.*?)\)
    # 2. Bold: \*\*(.*?)\*\*
    # 3. Italic: \*(.*?)\*
    # 4. Code: `(.*?)`
    
    pattern = r'(\[.*?\]\(.*?\)|`.*?`|\*\*.*?\*\*|\*.*?\*)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
        
        # Check link
        link_match = re.match(r'\[(.*?)\]\((.*?)\)', token)
        if link_match:
            link_text = link_match.group(1)
            run = paragraph.add_run(link_text)
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(47, 85, 151)
            run.font.underline = True
            if italic_default:
                run.italic = True
            continue
            
        # Check bold
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.bold = True
            run.font.color.rgb = RGBColor(38, 38, 38)
            if italic_default:
                run.italic = True
            continue
            
        # Check italic
        if token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.italic = True
            run.font.color.rgb = RGBColor(38, 38, 38)
            continue
            
        # Check inline code
        if token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 40, 40)
            continue
            
        # Normal text
        run = paragraph.add_run(token)
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(38, 38, 38)
        if italic_default:
            run.italic = True

def add_header_footer(doc, title_text):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"DE Genesis | {title_text}")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Trang | Dự Án DE Genesis - Báo Cáo Mentor")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def add_cover_page(doc, title_text):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # 1. Top Decorative Header Banner
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    cell = header_table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, PRIMARY_NAVY)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    hp = cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hrun = hp.add_run("DE GENESIS  |  DATA ENGINEERING ROADMAP")
    hrun.font.name = "Calibri"
    hrun.font.size = Pt(11)
    hrun.font.bold = True
    hrun.font.color.rgb = RGBColor(255, 255, 255)

    # 2. Main Title & Subtitle
    t_para = doc.add_paragraph()
    t_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_para.paragraph_format.space_before = Pt(20)
    t_para.paragraph_format.space_after = Pt(8)
    r_title = t_para.add_run(title_text.upper())
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(26)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub.paragraph_format.space_after = Pt(16)
    r_sub = sub.add_run("Xây Dựng & Vận Hành Hệ Thống Dữ Liệu Thực Tế — Từ OLTP 3NF Đến Lakehouse & Real-time Streaming")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(47, 85, 151)
    r_sub.bold = True

    # Divider bar
    div_table = doc.add_table(rows=1, cols=1)
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_background(div_cell, SECONDARY_BLUE)
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(12)
    p_sp2.paragraph_format.space_after = Pt(4)

    # 3. Callout Box: System Pillars & Overview
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_table.autofit = False
    box_cell = box_table.cell(0, 0)
    box_cell.width = Inches(6.5)
    set_cell_background(box_cell, LIGHT_BG)
    set_cell_margins(box_cell, top=120, bottom=120, left=160, right=160)
    
    box_tcPr = box_cell._element.get_or_add_tcPr()
    box_borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{PRIMARY_NAVY}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    box_tcPr.append(box_borders)

    bp0 = box_cell.paragraphs[0]
    bp0.paragraph_format.space_before = Pt(0)
    bp0.paragraph_format.space_after = Pt(6)
    r_box_head = bp0.add_run("📌 TỔNG QUAN HỆ THỐNG & NĂNG LỰC CHÍNH")
    r_box_head.font.name = "Calibri"
    r_box_head.font.size = Pt(10.5)
    r_box_head.font.bold = True
    r_box_head.font.color.rgb = RGBColor(31, 78, 121)

    bullets = [
        "Nền Tảng Core: Python, Java, SQL Phân Tích / Tối Ưu Execution Plan, Linux CLI.",
        "Mô Hình Dữ Liệu: Chuẩn hóa OLTP 3NF & Thiết kế Kimball Star Schema OLAP (Olist E-Commerce).",
        "Xử Lý & Stream: Batch Data Processing trên PySpark / Iceberg & Real-time Kafka Streaming.",
        "Tự Động Hóa & Giám Sát: Pipeline Orchestration với Apache Airflow, Docker Containerization."
    ]
    for b in bullets:
        bp = box_cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.line_spacing = 1.15
        r_bullet = bp.add_run(f"•  {b}")
        r_bullet.font.name = "Calibri"
        r_bullet.font.size = Pt(9.5)
        r_bullet.font.color.rgb = RGBColor(38, 38, 38)

    p_sp3 = doc.add_paragraph()
    p_sp3.paragraph_format.space_before = Pt(14)
    p_sp3.paragraph_format.space_after = Pt(4)

    # 4. Metadata Table
    meta_data = [
        ("Dự Án / Chương Trình", "DE Genesis — Lộ Trình Thực Hành Data Engineering 6 Tuần"),
        ("Đối Tượng Báo Cáo", "Mentor & Hội Đồng Phản Biện Kỹ Thuật"),
        ("Đội Ngũ Thực Hiện", "Data Engineering Team (DE-Genesis)"),
        ("Bộ Dữ Liệu Thử Nghiệm", "Brazilian E-Commerce Public Dataset by Olist (100k+ Orders)"),
        ("Môi Trường Vận Hành", "Docker, PostgreSQL 16, MySQL 8.0, Apache Kafka, Apache Spark, Airflow"),
        ("Thời Gian Nghiệm Thu", "Tháng 08 / 2026"),
    ]

    meta_table = doc.add_table(rows=len(meta_data), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    for row_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        
        c0 = row.cells[0]
        c0.width = Inches(2.2)
        set_cell_background(c0, LIGHT_BG if row_idx % 2 == 0 else "E9EEF4")
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        r0.font.name = "Calibri"
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(31, 78, 121)

        c1 = row.cells[1]
        c1.width = Inches(4.3)
        set_cell_background(c1, "FFFFFF" if row_idx % 2 == 0 else ZEBRA_BG)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(val)
        r1.font.name = "Calibri"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(38, 38, 38)

    set_table_borders(meta_table, color="D3D3D3", sz="4")

    p_end = doc.add_paragraph()
    p_end.paragraph_format.space_before = Pt(20)
    doc.add_page_break()

def convert_md_to_docx(md_filepath, docx_filepath, title):
    doc = Document()
    add_header_footer(doc, title)
    add_cover_page(doc, title)
    
    with open(md_filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    in_blockquote = False
    quote_text = ""
    
    def process_table_lines(tbl_lines):
        if not tbl_lines:
            return
        
        parsed_rows = []
        for line in tbl_lines:
            line_str = line.strip()
            if not line_str.startswith("|"):
                continue
            # Split by |
            parts = [p.strip() for p in line_str.split("|")[1:-1]]
            # Check separator line (---)
            if all(re.match(r'^:?-+:?$', p) for p in parts if p):
                continue
            parsed_rows.append(parts)
            
        if not parsed_rows:
            return
            
        num_cols = max(len(r) for r in parsed_rows)
        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, color=BORDER_GRAY, sz="4")
        
        for row_idx, row_data in enumerate(parsed_rows):
            is_header = (row_idx == 0)
            row = table.rows[row_idx]
            bg_color = PRIMARY_NAVY if is_header else (ZEBRA_BG if row_idx % 2 == 1 else "FFFFFF")
            
            for col_idx in range(num_cols):
                cell = row.cells[col_idx]
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                
                text_content = row_data[col_idx] if col_idx < len(row_data) else ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                if is_header:
                    run = p.add_run(text_content)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    parse_formatted_text(p, text_content, font_size=9.5)
                    
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Check code block fences ```
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                # flush any pending table or blockquote
                if in_table:
                    process_table_lines(table_lines)
                    table_lines = []
                    in_table = False
                if in_blockquote:
                    add_callout(doc, quote_text)
                    quote_text = ""
                    in_blockquote = False
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line.rstrip("\r\n"))
            i += 1
            continue
            
        # Check markdown table lines
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                if in_blockquote:
                    add_callout(doc, quote_text)
                    quote_text = ""
                    in_blockquote = False
                in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                process_table_lines(table_lines)
                table_lines = []
                in_table = False

        # Check blockquotes (> text)
        if stripped.startswith(">"):
            q_content = stripped[1:].strip()
            if not in_blockquote:
                in_blockquote = True
                quote_text = q_content
            else:
                quote_text += " " + q_content
            i += 1
            continue
        else:
            if in_blockquote:
                add_callout(doc, quote_text)
                quote_text = ""
                in_blockquote = False

        # Blank line
        if not stripped:
            i += 1
            continue
            
        # Horizontal rule ---
        if stripped in ["---", "***", "___"]:
            # Add thin divider or extra space
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
            
        # Headings
        if stripped.startswith("# "):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[2:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
            i += 1
            continue
            
        if stripped.startswith("## "):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[3:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(47, 85, 151)
            i += 1
            continue
            
        if stripped.startswith("### "):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[4:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(12.5)
            run.bold = True
            run.font.color.rgb = RGBColor(38, 38, 38)
            i += 1
            continue

        if stripped.startswith("#### "):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(stripped[5:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(38, 38, 38)
            i += 1
            continue

        # Bullet lists (- item or * item)
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            parse_formatted_text(p, stripped[2:].strip())
            i += 1
            continue

        # Numbered lists (1. item)
        num_match = re.match(r'^\d+\.\s+(.*)$', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            parse_formatted_text(p, num_match.group(1).strip())
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        parse_formatted_text(p, stripped)
        i += 1

    # Flush remaining table or blockquote if any
    if in_table:
        process_table_lines(table_lines)
    if in_blockquote:
        add_callout(doc, quote_text)

    try:
        doc.save(docx_filepath)
        print(f"Successfully exported: {docx_filepath}")
    except PermissionError:
        alt_filepath = docx_filepath.replace(".docx", "_WITH_COVER.docx")
        doc.save(alt_filepath)
        print(f"[WARNING] File is locked in MS Word. Exported to alternative file: {alt_filepath}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    files_to_export = [
        ("BAO_CAO_TONG_THE_DU_AN.md", "BAO_CAO_CONG_VIEC_6_TUAN_DATA_ENGINEERING.docx", "Báo Cáo Công Việc 6 Tuần Data Engineering"),
    ]
    
    for md_name, docx_name, title in files_to_export:
        md_path = os.path.join(base_dir, md_name)
        docx_path = os.path.join(base_dir, docx_name)
        if os.path.exists(md_path):
            convert_md_to_docx(md_path, docx_path, title)

